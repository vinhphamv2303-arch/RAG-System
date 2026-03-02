import re
import json
import os
import glob
from docx import Document


def clean_text(text):
    if not text: return ""
    text = text.replace('\xa0', ' ').replace('\t', ' ')
    text = re.sub(r'\s+', ' ', text)
    return text.strip()


def parse_header_v6(raw_lines):
    lines = [clean_text(line) for line in raw_lines[:60] if clean_text(line)]
    full_text = "\n".join(lines)
    meta = {"issuing_body": "Không rõ", "doc_number": "Không rõ", "issue_date": "Không rõ", "doc_name": "Không rõ",
            "basis": []}

    num_m = re.search(r"(Số|Luật số|Nghị định số|Thông tư số)[:\s]*([0-9/A-Z\-]+)", full_text, re.I)
    if num_m: meta["doc_number"] = num_m.group(2)

    date_m = re.search(r"ngày\s*(\d+)\s*tháng\s*(\d+)\s*năm\s*(\d+)", full_text, re.I)
    if date_m:
        meta["issue_date"] = f"{date_m.group(3)}-{date_m.group(2).zfill(2)}-{date_m.group(1).zfill(2)}"
    else:
        date_m2 = re.search(r"(\d{1,2})[/.-](\d{1,2})[/.-](\d{4})", full_text)
        if date_m2: meta["issue_date"] = f"{date_m2.group(3)}-{date_m2.group(2).zfill(2)}-{date_m2.group(1).zfill(2)}"

    for i, line in enumerate(lines):
        if line in ["LUẬT", "NGHỊ ĐỊNH", "THÔNG TƯ", "QUYẾT ĐỊNH"] and i + 1 < len(lines):
            meta["doc_name"] = f"{line} {lines[i + 1]}"
            break
        elif any(k in line for k in ["LUẬT ", "NGHỊ ĐỊNH ", "THÔNG TƯ "]) and line.isupper():
            meta["doc_name"] = line
            break

    for line in lines:
        if line.isupper() and len(line) < 50 and "CỘNG HÒA" not in line and line not in ["LUẬT", "NGHỊ ĐỊNH",
                                                                                         "THÔNG TƯ"]:
            clean_body = re.sub(r'[_+\-]', '', line).strip()
            if clean_body:
                meta["issuing_body"] = clean_body
                break

    basis_lines = []
    is_basis_section = False
    for line in lines:
        if line.startswith("Căn cứ"):
            is_basis_section = True
            basis_lines.append(line)
        elif is_basis_section:
            if re.match(r"^(Chương\s+[IVXLCDM]+|Điều\s+\d+)", line, re.IGNORECASE):
                break
            basis_lines.append(line)

    meta["basis"] = basis_lines
    return meta


def process_single_file(docx_path, output_dir):
    """Hàm xử lý cho 1 file Word cụ thể"""
    doc = Document(docx_path)
    header_lines = []
    for table in doc.tables[:3]:
        for row in table.rows:
            for cell in row.cells: header_lines.append(cell.text)
    header_lines.extend([p.text for p in doc.paragraphs[:60]])

    meta = parse_header_v6(header_lines)

    # Logic đặt tên file: Nếu tìm được số hiệu thì lấy số hiệu, nếu Không rõ thì lấy tên file gốc
    if meta["doc_number"] != "Không rõ":
        doc_prefix = meta["doc_number"].replace("/", "_")
    else:
        # Lấy tên file gốc (bỏ đuôi .docx) làm ID dự phòng để không bị ghi đè
        base_name = os.path.basename(docx_path)
        doc_prefix = os.path.splitext(base_name)[0].replace(" ", "_")

    final_json_structure = {
        "document_info": meta,
        "chunks": []
    }

    cur = {"chuong": "", "muc": "", "dieu_id": "", "dieu_full": "", "khoan_id": ""}
    is_body = False

    re_chuong_label = r"^Chương\s+[IVXLCDM]+$"
    re_muc_label = r"^Mục\s+\d+$"
    re_dieu = r"^(Điều\s+(\d+[a-zA-Z]*))\.(.*)"
    re_khoan = r"^(\d+)\.(.*)"
    re_diem = r"^([a-zđ])\)(.*)"

    paragraphs = doc.paragraphs
    for i, para in enumerate(paragraphs):
        line = clean_text(para.text)
        if not line: continue

        if re.match(re_chuong_label, line, re.I):
            title = clean_text(paragraphs[i + 1].text) if i + 1 < len(paragraphs) else ""
            cur["chuong"] = f"{line} - {title}"
            cur["muc"] = cur["dieu_id"] = cur["khoan_id"] = ""
            continue

        if re.match(re_muc_label, line, re.I):
            title = clean_text(paragraphs[i + 1].text) if i + 1 < len(paragraphs) else ""
            cur["muc"] = f"{line} - {title}"
            cur["dieu_id"] = cur["khoan_id"] = ""
            continue

        m_dieu = re.match(re_dieu, line, re.I)
        if m_dieu:
            is_body = True
            cur["dieu_id"] = m_dieu.group(1).replace(" ", "")
            cur["dieu_full"] = line
            cur["khoan_id"] = ""
            final_json_structure["chunks"].append({
                "chunk_id": f"{doc_prefix}_{cur['dieu_id']}",
                "hierarchy": {"chuong": cur["chuong"], "muc": cur["muc"], "dieu": cur["dieu_full"]},
                "content": line
            })
            continue

        if not is_body: continue

        m_khoan = re.match(re_khoan, line)
        if m_khoan and cur["dieu_id"]:
            cur["khoan_id"] = f"K{m_khoan.group(1)}"
            final_json_structure["chunks"].append({
                "chunk_id": f"{doc_prefix}_{cur['dieu_id']}_{cur['khoan_id']}",
                "hierarchy": {"chuong": cur["chuong"], "muc": cur["muc"], "dieu": cur["dieu_full"], "khoan": line},
                "content": line
            })
            continue

        m_diem = re.match(re_diem, line)
        if m_diem and cur["khoan_id"]:
            d_char = m_diem.group(1)
            final_json_structure["chunks"].append({
                "chunk_id": f"{doc_prefix}_{cur['dieu_id']}_{cur['khoan_id']}_Diem_{d_char}",
                "hierarchy": {"chuong": cur["chuong"], "muc": cur["muc"], "dieu": cur["dieu_full"],
                              "khoan": cur["khoan_id"], "diem": line},
                "content": line
            })
            continue

        if final_json_structure["chunks"] and is_body:
            if not re.match(re_chuong_label, line, re.I) and not re.match(re_muc_label, line, re.I):
                final_json_structure["chunks"][-1]["content"] += f"\n{line}"

    # Đảm bảo đường dẫn lưu file chính xác
    output_filename = os.path.join(output_dir, f"{doc_prefix}.json")
    with open(output_filename, "w", encoding="utf-8") as f:
        json.dump(final_json_structure, f, ensure_ascii=False, indent=4)

    return output_filename


def run_batch_pipeline(input_dir, output_dir):
    """Hàm chạy quét hàng loạt toàn bộ thư mục"""
    # 1. Tạo thư mục đích nếu nó chưa tồn tại
    if not os.path.exists(output_dir):
        os.makedirs(output_dir)
        print(f"📁 Đã tạo thư mục: {output_dir}")

    # 2. Quét tìm tất cả các file .docx
    search_pattern = os.path.join(input_dir, "*.docx")
    docx_files = glob.glob(search_pattern)

    if not docx_files:
        print(f"⚠️ Không tìm thấy file .docx nào trong thư mục '{input_dir}'")
        return

    print(f"🚀 Bắt đầu xử lý hàng loạt {len(docx_files)} file...\n" + "-" * 40)

    success_count = 0
    # 3. Vòng lặp xử lý từng file
    for file_path in docx_files:
        file_name = os.path.basename(file_path)
        print(f"⏳ Đang đọc: {file_name}")

        try:
            saved_path = process_single_file(file_path, output_dir)
            print(f"   ✅ Đã lưu: {os.path.basename(saved_path)}")
            success_count += 1
        except Exception as e:
            # Nếu có lỗi, in ra và đi tiếp sang file sau
            print(f"   ❌ LỖI bỏ qua file: {e}")

    print("-" * 40)
    print(f"🎉 Hoàn tất! Đã xử lý thành công {success_count}/{len(docx_files)} file.")


# ==========================================
# MAIN
# ==========================================
if __name__ == "__main__":
    RAW_DIR = "/home/vinh/projects/rag-system/data/raw"
    PREPROCESSED_DIR = "/home/vinh/projects/rag-system/data/preprocessed"

    run_batch_pipeline(RAW_DIR, PREPROCESSED_DIR)